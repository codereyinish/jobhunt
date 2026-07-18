"""Assisted ATS auto-fill.

Opens a *visible* Chromium on a job's application form, fills the fields it can
recognize from your profile (name, contact, links, and the work-authorization
questions that matter most for F-1 OPT), then hands YOU the browser to review
and click Submit. Nothing is ever submitted automatically.

Run:  python -m jobhunt.apply.autofill <job_id>
      python -m jobhunt.apply.autofill --url <application_url>   # ad-hoc test
"""
from __future__ import annotations

import sys

from ..core import db
from ..core.config import profile

# ── What to fill. Each concept: the value, and label patterns to recognize it.
#    Order matters — first matching rule wins for a given field.


def _profile_values() -> dict:
    p = profile()
    name = (p.get("name") or "").strip()
    first, _, last = name.partition(" ")
    wa = p.get("work_authorization") or {}
    needs = bool(wa.get("needs_sponsorship"))
    return {
        "first": first,
        "last": last or first,
        "name": name,
        "email": (p.get("email") or "").strip(),
        "phone": (p.get("phone") or "").strip(),
        "linkedin": (p.get("linkedin") or "").strip(),
        "github": (p.get("github") or "").strip(),
        "portfolio": (p.get("portfolio") or "").strip(),
        "location": (p.get("location") or "").strip(),
        "street": (p.get("address_line1") or p.get("address") or "").strip(),
        "line2": (p.get("address_line2") or "").strip(),
        "city": (p.get("city") or "").strip(),
        "state": (p.get("state") or "").strip(),
        "zip": (p.get("zip") or p.get("postal_code") or "").strip(),
        "country": (p.get("country") or "United States").strip(),
        "resume_path": (p.get("resume_path") or "").strip(),
        # yes = authorized to work now (OPT); needs sponsorship in the future.
        "authorized": "yes",
        "sponsorship": "yes" if needs else "no",
    }


def _work_experience() -> list:
    """Structured work-experience entries from profile.yaml, for the repeatable
    'Work Experience' blocks on application forms."""
    out = []
    for e in (profile().get("work_experience") or []):
        out.append({
            "title": (e.get("title") or "").strip(),
            "company": (e.get("company") or "").strip(),
            "location": (e.get("location") or "").strip(),
            "start": (e.get("start") or "").strip(),
            "end": (e.get("end") or "").strip(),
            "current": bool(e.get("current")),
            "description": " ".join((e.get("description") or "").split()),
        })
    return out


# concept -> label regexes (matched against label text, lowercased). ORDER MATTERS —
# first match wins, so the specific address rules sit above the generic location/name.
_TEXT_RULES = [
    ("email", r"e-?mail"),
    ("phone", r"phone|mobile|cell"),
    ("linkedin", r"linkedin"),
    ("github", r"github"),
    ("portfolio", r"portfolio|website|personal site|personal url"),
    ("line2", r"address line ?2|address 2|\bapt\b|\bsuite\b|\bunit\b|apartment"),
    ("street", r"address line ?1|address 1|street address|mailing address|^\s*address\b|\bstreet\b"),
    ("zip", r"postal code|zip ?code|\bzip\b|postcode|post code"),
    ("city", r"\bcity\b|\btown\b"),
    ("state", r"\bstate\b|province|region"),
    ("country", r"\bcountry\b"),
    ("first", r"first name|given name|legal first"),
    ("last", r"last name|surname|family name|legal last"),
    ("location", r"where do you (live|reside)|current location|residence|your location"),
    ("name", r"^\s*(full |your )?name\b|full legal name"),
]

# US state name ↔ abbreviation, so a State <select> matches whether your profile
# stores "New York" or "NY", and whichever the form's options use.
_STATE_FULL = {
    "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas", "CA": "California",
    "CO": "Colorado", "CT": "Connecticut", "DE": "Delaware", "FL": "Florida", "GA": "Georgia",
    "HI": "Hawaii", "ID": "Idaho", "IL": "Illinois", "IN": "Indiana", "IA": "Iowa",
    "KS": "Kansas", "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine", "MD": "Maryland",
    "MA": "Massachusetts", "MI": "Michigan", "MN": "Minnesota", "MS": "Mississippi",
    "MO": "Missouri", "MT": "Montana", "NE": "Nebraska", "NV": "Nevada", "NH": "New Hampshire",
    "NJ": "New Jersey", "NM": "New Mexico", "NY": "New York", "NC": "North Carolina",
    "ND": "North Dakota", "OH": "Ohio", "OK": "Oklahoma", "OR": "Oregon", "PA": "Pennsylvania",
    "RI": "Rhode Island", "SC": "South Carolina", "SD": "South Dakota", "TN": "Tennessee",
    "TX": "Texas", "UT": "Utah", "VT": "Vermont", "VA": "Virginia", "WA": "Washington",
    "WV": "West Virginia", "WI": "Wisconsin", "WY": "Wyoming", "DC": "Washington DC",
}

# Never treat these as fillable text — they're custom dropdowns/essays where a
# wrong free-text answer does damage. If a label hits one, we leave it for you.
_SKIP = r"relocat|open to work|earliest|deadline|timeline|why |cover letter|" \
        r"salary|compensation|how did you hear|referr|pronoun|gender|race|" \
        r"veteran|disability|describe|tell us|what interests"

# select/radio questions -> which stored yes/no answer to use
_CHOICE_RULES = [
    ("sponsorship", r"sponsor|visa sponsor|require.*visa|now or in the future"),
    ("authorized", r"authorized to work|legally authorized|eligible to work|work authorization"),
]


def _match(label: str, rules) -> str | None:
    import re
    for concept, pat in rules:
        if re.search(pat, label):
            return concept
    return None


def _skippable(label: str) -> bool:
    import re
    return bool(re.search(_SKIP, label))


_JS_COLLECT = r"""() => {
  const els = [...document.querySelectorAll('input, textarea, select')];
  const out = [];
  els.forEach((el) => {
    const type = (el.type || '').toLowerCase();
    if (['hidden', 'submit', 'button', 'reset', 'image', 'checkbox', 'radio'].includes(type)) return;
    const off = el.offsetParent === null && el.tagName !== 'SELECT';
    if (off) return;                                   // skip hidden fields
    const idx = out.length;
    el.setAttribute('data-jh', idx);
    let label = '';
    if (el.id) {
      try { const l = document.querySelector('label[for="' + CSS.escape(el.id) + '"]');
            if (l) label += ' ' + l.innerText; } catch (e) {}
    }
    const anc = el.closest('label, .field, .form-group, .application-field, [class*=field]');
    if (anc) label += ' ' + (anc.innerText || '');
    label += ' ' + (el.getAttribute('aria-label') || '') + ' ' +
             (el.name || '') + ' ' + (el.placeholder || '');
    out.push({
      idx, tag: el.tagName.toLowerCase(), type,
      label: label.replace(/\s+/g, ' ').trim().toLowerCase().slice(0, 200),
      options: el.tagName === 'SELECT' ? [...el.options].map(o => o.text) : []
    });
  });
  // resume/CV file inputs, tracked separately
  const files = [];
  document.querySelectorAll('input[type=file]').forEach((el, i) => {
    el.setAttribute('data-jhf', i);
    const anc = el.closest('label, .field, [class*=field]');
    const lab = ((anc ? anc.innerText : '') + ' ' + (el.name || '') + ' ' +
                 (el.getAttribute('aria-label') || '')).toLowerCase();
    files.push({ i, label: lab });
  });
  return { fields: out, files };
}"""


def _choose_option(options: list[str], want: str) -> str | None:
    """Pick the option text matching a yes/no intent."""
    want = want.lower()
    for o in options:
        t = o.strip().lower()
        if want == "yes" and t in ("yes", "yes, i am", "yes, i do") or t == want:
            return o
    for o in options:                                  # loose: startswith
        if o.strip().lower().startswith(want):
            return o
    return None


def _choose_named(options: list[str], want: str) -> str | None:
    """Pick a dropdown option by name — handles State (full name ↔ abbreviation)
    and Country, regardless of which form the profile or the form uses."""
    want = (want or "").strip()
    if not want:
        return None
    cands = {want.lower()}
    if want.upper() in _STATE_FULL:                    # profile has "NY" → also try "New York"
        cands.add(_STATE_FULL[want.upper()].lower())
    for ab, full in _STATE_FULL.items():               # profile has "New York" → also try "NY"
        if full.lower() == want.lower():
            cands.add(ab.lower())
    for o in options:
        if o.strip().lower() in cands:
            return o
    for o in options:                                  # loose: startswith any candidate
        ol = o.strip().lower()
        if any(ol.startswith(c) for c in cands):
            return o
    return None


def _plan(data: dict, vals: dict) -> tuple[list, list]:
    """Pure decision step: from collected fields decide what to fill. Returns
    (actions, log) where each action is (kind, selector, arg). Shared by the sync
    and async executors so the matching logic never drifts."""
    actions, log = [], []
    for f in data["fields"]:
        sel = f"[data-jh='{f['idx']}']"
        label = f["label"]
        if _skippable(label):
            continue
        if f["tag"] == "select":
            concept = _match(label, _CHOICE_RULES) or _match(label, _TEXT_RULES)
            if concept in ("authorized", "sponsorship"):
                opt = _choose_option(f["options"], vals.get(concept, ""))
            elif concept in ("state", "country"):
                opt = _choose_named(f["options"], vals.get(concept, ""))
            else:
                continue
            if opt:
                actions.append(("select", sel, opt))
                log.append(f"  ✓ [{label[:40]}] → {opt}")
            continue
        concept = _match(label, _TEXT_RULES)
        val = vals.get(concept, "") if concept else ""
        if not val:
            continue
        actions.append(("fill", sel, val))
        log.append(f"  ✓ [{label[:40]}] → {val}")

    resume = vals.get("resume_path")
    if resume and data["files"]:
        target = next((f for f in data["files"] if "resume" in f["label"] or "cv" in f["label"]),
                      data["files"][0])
        actions.append(("file", f"[data-jhf='{target['i']}']", resume))
        log.append(f"  ✓ resume uploaded ({resume})")
    elif not resume:
        log.append("  ⚠ no resume_path in profile.yaml — upload it yourself")
    return actions, log


def fill(page, vals: dict) -> list[str]:
    """Sync filler (CLI / ad-hoc use). Fills everything recognized; returns a log."""
    actions, log = _plan(page.evaluate(_JS_COLLECT), vals)
    for kind, sel, arg in actions:
        try:
            if kind == "fill":
                page.fill(sel, arg)
            elif kind == "select":
                page.select_option(sel, label=arg)
            elif kind == "file":
                page.set_input_files(sel, arg)
        except Exception:
            pass
    return log


async def _afill(page, vals: dict) -> list[str]:
    """Async filler — safe to call from an exposed-binding callback."""
    actions, log = _plan(await page.evaluate(_JS_COLLECT), vals)
    for kind, sel, arg in actions:
        try:
            if kind == "fill":
                await page.fill(sel, arg)
            elif kind == "select":
                await page.select_option(sel, label=arg)
            elif kind == "file":
                await page.set_input_files(sel, arg)
        except Exception:
            pass
    return log


def _skill_match(job: dict) -> dict:
    """A lightweight Simplify-style 'resume match': how many of your profile skills
    actually appear in this job's description."""
    from ..core.config import profile
    skills = [s for s in (profile().get("skills") or []) if s]
    jd = (job.get("description") or "").lower()
    if not skills or not jd:
        return {"pct": 0, "have": 0, "total": len(skills)}
    have = sum(1 for s in skills if s.lower() in jd)
    return {"pct": round(100 * have / len(skills)), "have": have, "total": len(skills)}


def _cover_letter(job: dict, timeout: int = 180) -> str:
    """Draft a tailored cover letter (greeting → sign-off) from the FULL resume +
    JD via the claude CLI, in the candidate's established voice/structure."""
    from ..core.config import profile, resume_text
    from ..match.llm import _run, claude_available
    if not claude_available():
        return ""
    p = profile()
    resume = resume_text()
    company = job.get("company", "") or "the company"
    prompt = (
        "You are writing a cover letter for the candidate below, applying to the role below. "
        "Match this exact structure and voice — warm, specific, confident, first person, no "
        "clichés and no 'I am writing to express my interest':\n\n"
        "  • Start with 'Dear Hiring Team,'\n"
        "  • Paragraph 1 — who they are (a new-grad engineer, B.S. Math/CS, ~3 years building "
        "and shipping backend systems end to end) and their headline project as a hook.\n"
        "  • Paragraph 2 — go deep on ONE project with CONCRETE metrics from the resume "
        "(e.g. the ~30%/~65% ClassRec numbers) and explicitly connect that experience to what "
        f"THIS {company} role/team actually does.\n"
        "  • Paragraph 3 — breadth: 2-3 OTHER real projects from the resume, and their comfort "
        "working WITH AI coding tools day to day (not around them). Show genuine curiosity.\n"
        "  • Paragraph 4 — what they're looking for now (real ownership early, mentorship, room "
        "to grow into meaningful impact) and a one-line ask to talk.\n"
        "  • End with 'Thank you for your time and consideration.' then 'Sincerely,' on its own "
        "line, then the candidate's name on the final line.\n\n"
        "Ground EVERY claim in the resume — invent nothing, no placeholders, use the real "
        "company and role. ~320-400 words. Output ONLY the letter text from the greeting to the "
        "name — no letterhead, no date, no commentary.\n\n"
        f"CANDIDATE NAME: {p.get('name', '')}\n"
        f"ROLE: {job.get('title', '')} at {company}\n\n"
        f"JOB DESCRIPTION:\n{(job.get('description') or '')[:3500]}\n\n"
        f"RESUME (use these real projects + metrics):\n{resume}\n"
    )
    try:
        return _run(prompt, timeout).strip()
    except Exception:
        return ""


def _build_cover_docx(job: dict, body: str) -> str:
    """Assemble a .docx cover letter matching the candidate's format: name header,
    contact line, date, recipient + Re line, then the letter body. Returns the path."""
    import datetime
    import re as _re

    from docx import Document
    from docx.shared import Pt

    from ..core.config import DATA_DIR, profile
    p = profile()
    name = p.get("name", "") or "Applicant"
    contact = "  |  ".join(x for x in [
        p.get("location", ""), p.get("phone", ""), p.get("email", ""),
        (p.get("linkedin", "") or "").replace("https://", "").replace("http://", ""),
        (p.get("github", "") or "").replace("https://", "").replace("http://", ""),
    ] if x)

    doc = Document()
    hdr = doc.add_paragraph()
    run = hdr.add_run(name); run.bold = True; run.font.size = Pt(16)
    doc.add_paragraph(contact).runs[0].font.size = Pt(9)
    doc.add_paragraph()
    doc.add_paragraph(datetime.date.today().strftime("%B %-d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("Hiring Team")
    if job.get("company"):
        doc.add_paragraph(job["company"])
    if job.get("title"):
        re_p = doc.add_paragraph(); re_p.add_run(f"Re: {job['title']}").bold = True
    doc.add_paragraph()
    for line in (body or "").split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    out_dir = DATA_DIR / "cover_letters"
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = _re.sub(r"[^A-Za-z0-9]+", "_", f"{name}_Cover_Letter_{job.get('company','')}").strip("_")
    path = str(out_dir / f"{safe}.docx")
    doc.save(path)
    return path


def _open_file(path: str) -> None:
    import subprocess
    import sys
    try:
        if sys.platform == "darwin":
            subprocess.Popen(["open", path])
        elif sys.platform.startswith("win"):
            import os
            os.startfile(path)                       # type: ignore[attr-defined]
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception:
        pass


# Open-ended questions we CAN draft (essays), and ones we must never AI-answer
# (demographic/EEO + the plain profile fields autofill already handles).
_ANSWERABLE = (r"why |describe|tell us|what interests|what excites|how would|"
               r"your experience|motivat|passion|proud|challenge|interest in|"
               r"want to work|excited|what draws|why do you|why are you")
_NOT_ANSWERABLE = (r"gender|race|veteran|disabilit|pronoun|ethnic|hispanic|salary|"
                   r"compensation|how did you hear|referr|relocat|authorized|sponsor|"
                   r"cover letter|first name|last name|e-?mail|phone|linkedin|github")


def _answerable(label: str) -> bool:
    import re
    return bool(re.search(_ANSWERABLE, label)) and not re.search(_NOT_ANSWERABLE, label)


def _answer_question(job: dict, question: str, hint: str = "", timeout: int = 150) -> str:
    """Draft an answer to ONE application question from resume + JD (+ your hint)."""
    from ..core.config import profile, resume_text
    from ..match.llm import _run, claude_available
    if not claude_available():
        return ""
    p = profile()
    resume = resume_text()
    prompt = (
        "Answer this job-application question for the candidate, in the first person, ready to "
        "paste. Ground every claim in the RESUME — invent nothing. Be specific and concise "
        "(80-150 words unless the question implies otherwise). No clichés, do not restate the "
        "question.\n\n"
        f"QUESTION:\n{question}\n\n"
        + (f"CANDIDATE'S GUIDANCE (follow this closely):\n{hint}\n\n" if hint.strip() else "")
        + f"CANDIDATE NAME: {p.get('name', '')}\n"
        f"ROLE: {job.get('title', '')} at {job.get('company', '')}\n\n"
        f"JOB DESCRIPTION:\n{(job.get('description') or '')[:2000]}\n\n"
        f"RESUME:\n{resume[:4500]}\n"
    )
    try:
        return _run(prompt, timeout).strip()
    except Exception:
        return ""


_INSERT_JS = r"""(a) => {
  const el = document.querySelector(`[data-jh='${a.idx}']`);
  if (!el) return false;
  el.focus(); el.value = a.text;
  el.dispatchEvent(new Event('input', {bubbles: true}));
  el.dispatchEvent(new Event('change', {bubbles: true}));
  return true;
}"""


_COVER_JS = r"""(text) => {
  const els = [...document.querySelectorAll('textarea, input[type=text]')];
  for (const el of els) {
    const anc = el.closest('label, .field, [class*=field]');
    const lab = ((anc ? anc.innerText : '') + ' ' + (el.name || '') + ' ' +
                 (el.getAttribute('aria-label') || '') + ' ' + (el.placeholder || '')).toLowerCase();
    if (lab.includes('cover')) {
      el.focus(); el.value = text;
      el.dispatchEvent(new Event('input', {bubbles: true}));
      el.dispatchEvent(new Event('change', {bubbles: true}));
      return true;
    }
  }
  return false;
}"""


async def _afill_cover_field(page, text: str) -> bool:
    """Drop the cover letter into a cover-letter textarea on the page, if one exists."""
    try:
        return bool(await page.evaluate(_COVER_JS, text))
    except Exception:
        return False


# The floating Simplify-style panel, injected into the page. All state/spinners
# live in JS; buttons call the Python bindings exposed below.
_PANEL_JS = r"""(data) => {
  if (document.getElementById('jh-panel')) return;
  const style = document.createElement('style');
  style.id = 'jh-style';
  style.textContent = `
  #jh-panel, #jh-panel * { box-sizing: border-box; font-family: -apple-system,
    BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; }
  #jh-panel { position: fixed; top: 18px; right: 18px; width: 340px; z-index: 2147483647;
    background: #fff; border: 1px solid #e6e9ef; border-radius: 16px;
    box-shadow: 0 18px 60px rgba(20,30,60,.22); overflow: hidden; color: #1a2233; }
  .jh-head { display: flex; align-items: center; justify-content: space-between;
    padding: 13px 16px; border-bottom: 1px solid #eef1f6; }
  .jh-brand { display: flex; align-items: center; gap: 7px; font-weight: 700; font-size: 15px; }
  .jh-bolt { display: inline-flex; align-items: center; justify-content: center; width: 22px;
    height: 22px; background: linear-gradient(135deg,#3b82f6,#22d3c5); border-radius: 6px;
    color: #fff; font-size: 12px; }
  .jh-x { cursor: pointer; color: #9aa4b5; font-size: 21px; line-height: 1; padding: 0 4px; }
  .jh-x:hover { color: #556; }
  .jh-body { padding: 14px 16px 16px; }
  .jh-jobline { font-size: 12px; color: #6b7688; margin-bottom: 12px; white-space: nowrap;
    overflow: hidden; text-overflow: ellipsis; }
  .jh-match { display: flex; align-items: center; gap: 12px; background: #f5f8ff;
    border: 1px solid #e6edfb; border-radius: 12px; padding: 11px 13px; margin-bottom: 14px; }
  .jh-ring { flex-shrink: 0; width: 46px; height: 46px; border-radius: 50%; display: flex;
    align-items: center; justify-content: center; font-size: 12px; font-weight: 700; color: #2563eb; }
  .jh-matchtxt { font-size: 12.5px; line-height: 1.4; color: #3d4757; }
  .jh-matchtxt b { color: #1a2233; }
  .jh-btn { width: 100%; border: none; border-radius: 10px; padding: 11px 14px; font-size: 14px;
    font-weight: 650; cursor: pointer; display: flex; align-items: center; justify-content: center;
    gap: 7px; transition: filter .12s, background .12s; }
  .jh-btn:disabled { opacity: .6; cursor: default; }
  .jh-primary { background: linear-gradient(135deg,#3b82f6,#2f6ef0); color: #fff; }
  .jh-primary:hover:not(:disabled) { filter: brightness(1.06); }
  .jh-outline { background: #fff; color: #2563eb; border: 1.5px solid #cfe0fd; }
  .jh-outline:hover:not(:disabled) { background: #f5f8ff; }
  .jh-done-btn { background: #ecfdf5; color: #059669; border: 1.5px solid #b6ebd4; }
  .jh-done-btn.done { background: #059669; color: #fff; border-color: #059669; }
  .jh-stat { font-size: 12.5px; color: #059669; font-weight: 600; margin-top: 8px;
    min-height: 0; text-align: center; }
  .jh-sep { height: 1px; background: #eef1f6; margin: 14px 0; }
  .jh-label { font-size: 10.5px; text-transform: uppercase; letter-spacing: .08em;
    color: #94a0b3; font-weight: 700; margin-bottom: 9px; }
  .jh-coverwrap { display: none; margin-top: 10px; }
  .jh-cover { width: 100%; border: 1px solid #e0e5ee; border-radius: 9px; padding: 9px 11px;
    font-size: 12.5px; line-height: 1.5; color: #1a2233; resize: vertical; font-family: inherit; }
  .jh-cover:focus { outline: none; border-color: #3b82f6; }
  .jh-row { display: flex; gap: 8px; margin-top: 8px; }
  .jh-mini { flex: 1; border: 1px solid #e0e5ee; background: #fff; border-radius: 8px;
    padding: 7px; font-size: 12px; font-weight: 600; color: #3d4757; cursor: pointer; }
  .jh-mini:hover { background: #f5f7fb; }
  .jh-foot { font-size: 11px; color: #9aa4b5; line-height: 1.4; margin-top: 12px; text-align: center; }
  #jh-qlist { margin-top: 10px; display: flex; flex-direction: column; gap: 10px; }
  .jh-qempty { font-size: 12px; color: #9aa4b5; text-align: center; padding: 6px 0; }
  .jh-q { border: 1px solid #eef1f6; border-radius: 10px; padding: 10px 11px; background: #fbfcfe; }
  .jh-qlabel { font-size: 12px; font-weight: 600; color: #3d4757; line-height: 1.35; margin-bottom: 8px; }
  .jh-qhint { width: 100%; border: 1px solid #e0e5ee; border-radius: 8px; padding: 7px 9px;
    font-size: 12px; color: #1a2233; margin-bottom: 8px; font-family: inherit; }
  .jh-qhint:focus { outline: none; border-color: #3b82f6; }
  .jh-qgen { width: 100%; }
  #jh-panel .jh-cover { max-height: 220px; }
  `;
  document.head.appendChild(style);

  const pct = data.match.total ? data.match.pct : 0;
  const wrap = document.createElement('div'); wrap.id = 'jh-panel';
  wrap.innerHTML = `
    <div class="jh-head">
      <div class="jh-brand"><span class="jh-bolt">&#9889;</span> jobhunt</div>
      <span id="jh-close" class="jh-x">&times;</span>
    </div>
    <div class="jh-body">
      <div class="jh-jobline">${data.title || ''}${data.company ? ' &middot; ' + data.company : ''}</div>
      <div class="jh-match">
        <div class="jh-ring" id="jh-ring">${pct}%</div>
        <div class="jh-matchtxt"><b>${data.match.have} of ${data.match.total}</b> of your skills appear in this job</div>
      </div>
      <button id="jh-autofill" class="jh-btn jh-primary">&#9889; Autofill this page</button>
      <div id="jh-fillstat" class="jh-stat"></div>
      <div class="jh-sep"></div>
      <div class="jh-label">Cover letter</div>
      <button id="jh-cover" class="jh-btn jh-outline">&#9998; Generate with AI</button>
      <div id="jh-coverwrap" class="jh-coverwrap">
        <textarea id="jh-covertext" class="jh-cover" rows="8" placeholder="Your cover letter will appear here…"></textarea>
        <div class="jh-row">
          <button id="jh-copy" class="jh-mini">Copy</button>
          <button id="jh-insert" class="jh-mini">Insert into form</button>
        </div>
        <button id="jh-docx" class="jh-btn jh-outline" style="margin-top:8px">&#128196; Save &amp; open .docx</button>
      </div>
      <div class="jh-sep"></div>
      <div class="jh-label">Application questions</div>
      <button id="jh-scan" class="jh-btn jh-outline">&#9906; Scan this page for questions</button>
      <div id="jh-qlist"></div>
      <div class="jh-sep"></div>
      <div class="jh-label">Work experience</div>
      <button id="jh-exp" class="jh-btn jh-outline">&#128188; Fill work experience</button>
      <div id="jh-expstat" class="jh-stat"></div>
      <div class="jh-sep"></div>
      <button id="jh-applied" class="jh-btn jh-done-btn">Mark applied &rarr;</button>
      <div class="jh-foot">Review every field before you submit — nothing is submitted for you.</div>
    </div>`;
  document.body.appendChild(wrap);

  const ring = wrap.querySelector('#jh-ring');
  const col = pct >= 60 ? '#059669' : (pct >= 30 ? '#2563eb' : '#e0863a');
  ring.style.background = `conic-gradient(${col} ${pct * 3.6}deg, #e6edfb 0deg)`;
  ring.style.color = col;
  const inner = document.createElement('div');
  inner.style.cssText = 'width:36px;height:36px;border-radius:50%;background:#fff;display:flex;align-items:center;justify-content:center;';
  inner.textContent = pct + '%'; ring.textContent = ''; ring.appendChild(inner);

  const $ = s => wrap.querySelector(s);
  $('#jh-close').onclick = () => { wrap.remove(); };
  $('#jh-autofill').onclick = async (e) => {
    const b = e.currentTarget, o = b.innerHTML; b.disabled = true; b.textContent = 'Filling…';
    try { const r = await window.jhAutofill();
      $('#jh-fillstat').textContent = r && r.count ? ('✓ Filled ' + r.count + ' field' + (r.count == 1 ? '' : 's') + ' — review them')
                                                   : 'No recognizable fields — fill manually';
    } catch (err) { $('#jh-fillstat').textContent = 'Autofill failed'; }
    b.disabled = false; b.innerHTML = o;
  };
  $('#jh-cover').onclick = async (e) => {
    const b = e.currentTarget, o = b.innerHTML; b.disabled = true; b.textContent = 'Writing… ~30s';
    $('#jh-coverwrap').style.display = 'block';
    try { const t = await window.jhCover();
      $('#jh-covertext').value = t || 'Could not generate — is the claude CLI available?';
    } catch (err) { $('#jh-covertext').value = 'Could not generate the cover letter.'; }
    b.disabled = false; b.innerHTML = o;
  };
  $('#jh-copy').onclick = () => { const t = $('#jh-covertext'); t.select();
    try { document.execCommand('copy'); } catch (e) {}
    $('#jh-copy').textContent = 'Copied ✓'; setTimeout(() => $('#jh-copy').textContent = 'Copy', 1500); };
  $('#jh-insert').onclick = async () => {
    const ok = await window.jhInsertCover($('#jh-covertext').value);
    $('#jh-insert').textContent = ok ? 'Inserted ✓' : 'No field on page';
    setTimeout(() => $('#jh-insert').textContent = 'Insert into form', 1800); };
  $('#jh-docx').onclick = async (e) => {
    const b = e.currentTarget, o = b.innerHTML; b.disabled = true; b.textContent = 'Building .docx…';
    let name = ''; try { name = await window.jhCoverDocx($('#jh-covertext').value); } catch (err) {}
    b.textContent = name ? ('Opened ' + name + ' ✓') : 'Could not build .docx';
    setTimeout(() => { b.disabled = false; b.innerHTML = o; }, 3000); };
  $('#jh-applied').onclick = async (e) => {
    const b = e.currentTarget; b.disabled = true;
    try { await window.jhApplied(); } catch (err) {}
    b.classList.add('done'); b.textContent = 'Applied ✓ — added to tracker'; };

  $('#jh-scan').onclick = async (e) => {
    const b = e.currentTarget, o = b.innerHTML; b.disabled = true; b.textContent = 'Scanning…';
    let qs = []; try { qs = await window.jhQuestions(); } catch (err) {}
    const list = $('#jh-qlist'); list.innerHTML = '';
    if (!qs || !qs.length) {
      list.innerHTML = '<div class="jh-qempty">No open-ended questions found on this page.</div>';
    }
    (qs || []).forEach((q) => {
      const card = document.createElement('div'); card.className = 'jh-q';
      let raw = q.label || 'Question';
      const m = raw.match(/^(.*?[?.:])(\s|$)/);           // collector repeats the label — keep 1st sentence
      let lab = (m ? m[1] : raw).replace(/</g, '&lt;').slice(0, 140);
      card.innerHTML = `
        <div class="jh-qlabel">${lab}</div>
        <input class="jh-qhint" placeholder="Add a hint to steer the answer (optional)">
        <button class="jh-qgen jh-mini">&#10024; Generate answer</button>
        <textarea class="jh-qans jh-cover" rows="5" style="display:none;margin-top:8px"></textarea>
        <div class="jh-row jh-qrow" style="display:none">
          <button class="jh-qcopy jh-mini">Copy</button>
          <button class="jh-qins jh-mini">Insert into form</button>
        </div>`;
      const gen = card.querySelector('.jh-qgen'), ans = card.querySelector('.jh-qans'),
            row = card.querySelector('.jh-qrow'), hint = card.querySelector('.jh-qhint');
      gen.onclick = async () => {
        const oo = gen.innerHTML; gen.disabled = true; gen.textContent = 'Writing… ~30s';
        let t = ''; try { t = await window.jhAnswer(q.label, hint.value); } catch (err) {}
        ans.style.display = 'block'; row.style.display = 'flex';
        ans.value = t || 'Could not generate — is the claude CLI available?';
        gen.disabled = false; gen.innerHTML = oo;
      };
      card.querySelector('.jh-qcopy').onclick = () => { ans.select();
        try { document.execCommand('copy'); } catch (e) {} };
      card.querySelector('.jh-qins').onclick = async (ev) => {
        const ok = await window.jhInsertField(q.idx, ans.value);
        ev.currentTarget.textContent = ok ? 'Inserted ✓' : 'Failed';
        setTimeout(() => ev.currentTarget.textContent = 'Insert into form', 1600);
      };
      list.appendChild(card);
    });
    b.disabled = false; b.innerHTML = o;
  };

  // Fill repeatable "Work Experience" blocks from your saved entries.
  function _lbl(el){
    let s = '';
    if (el.id){ try { const l = document.querySelector('label[for="' + CSS.escape(el.id) + '"]');
      if (l) s += ' ' + l.innerText; } catch(e){} }
    const anc = el.closest('label, .field, .form-group, [class*=field]');
    if (anc) s += ' ' + (anc.innerText || '');
    s += ' ' + (el.getAttribute('aria-label') || '') + ' ' + (el.name || '') + ' ' + (el.placeholder || '');
    return s.replace(/\s+/g, ' ').trim().toLowerCase();
  }
  function _set(el, v){
    if (!el || v == null || v === '') return;
    el.focus(); el.value = v;
    el.dispatchEvent(new Event('input', {bubbles: true}));
    el.dispatchEvent(new Event('change', {bubbles: true}));
  }
  function _expBlocks(){
    const ctrls = [...document.querySelectorAll('input, textarea, select')]
      .filter(e => e.type !== 'hidden' && e.type !== 'file');
    const blocks = []; let cur = null;
    for (const el of ctrls){
      const L = _lbl(el);
      if (/job title|position title|role title|^title\b/.test(L)){ cur = {title: el, fields: []}; blocks.push(cur); }
      if (cur) cur.fields.push({el, L});
    }
    return blocks;
  }
  async function fillExperience(entries){
    if (!entries || !entries.length) return 0;
    const addBtn = () => [...document.querySelectorAll('button, a')]
      .find(b => /add another|add experience|add more|add position/i.test(b.textContent || ''));
    let n = 0;
    for (let i = 0; i < entries.length; i++){
      let blocks = _expBlocks();
      if (i >= blocks.length){ const a = addBtn(); if (a){ a.click(); await new Promise(r => setTimeout(r, 700)); blocks = _expBlocks(); } }
      const blk = blocks[i]; if (!blk) break;
      const e = entries[i];
      _set(blk.title, e.title);
      for (const {el, L} of blk.fields){
        if (el === blk.title) continue;
        if (/compan|employer|organi/.test(L)) _set(el, e.company);
        else if (/\blocation\b|\bcity\b/.test(L)) _set(el, e.location);
        else if (/from|start/.test(L)) _set(el, e.start);
        else if (/\bto\b|end date|\bend\b/.test(L)){ if (!e.current) _set(el, e.end); }
        else if (/role description|description|responsib|summary|about/.test(L)) _set(el, e.description);
      }
      if (e.current){
        const cb = [...document.querySelectorAll('input[type=checkbox]')]
          .find(x => /currently work here|current(ly)?|present|i work here/i.test(_lbl(x)));
        if (cb && !cb.checked) cb.click();
      }
      n++;
    }
    return n;
  }
  $('#jh-exp').onclick = async (e) => {
    const b = e.currentTarget, o = b.innerHTML; b.disabled = true; b.textContent = 'Filling…';
    let ents = []; try { ents = await window.jhExperience(); } catch (err) {}
    let n = 0; try { n = await fillExperience(ents); } catch (err) {}
    $('#jh-expstat').textContent = n
      ? ('✓ Filled ' + n + ' experience block' + (n == 1 ? '' : 's') + ' — set the dates + review')
      : 'No work experience saved — add it in profile.yaml';
    b.disabled = false; b.innerHTML = o;
  };
}"""


async def _launch(pw):
    """Pick the best browser to drive, in order of "most like your real browser":

    1. ATTACH to a Chrome you already started with --remote-debugging-port (run the
       chrome-debug helper) — the panel drops into your real, logged-in session.
    2. LAUNCH your Google Chrome with a dedicated persistent profile — logins persist
       across runs after a one-time sign-in.
    3. Fall back to the bundled Chromium.

    Returns (context, browser|None, attached). attached=True means we joined YOUR
    Chrome, so we open a fresh tab and must never close the browser on exit."""
    import os

    from ..core.config import DATA_DIR

    # 1) Attach over CDP if something is listening (you ran the chrome-debug helper).
    cdp = os.environ.get("JOBHUNT_CDP", "http://localhost:9222")
    try:
        browser = await pw.chromium.connect_over_cdp(cdp)
        ctx = browser.contexts[0] if browser.contexts else await browser.new_context()
        print(f"  ▶ attached to your running Chrome at {cdp}")
        return ctx, browser, True
    except Exception:
        pass                                           # nothing listening → launch our own

    # 2) Launch your Chrome with a dedicated persistent profile (login once, persists).
    profile_dir = DATA_DIR / "chrome-profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    try:
        ctx = await pw.chromium.launch_persistent_context(
            str(profile_dir), channel="chrome", headless=False,
            no_viewport=True, args=["--start-maximized"])
        print("  ▶ launched your Google Chrome (persistent jobhunt profile)")
        return ctx, None, False
    except Exception as e:
        print(f"  (couldn't open your Chrome — {str(e)[:80]}… — using the built-in browser)")
        browser = await pw.chromium.launch(headless=False, args=["--start-maximized"])
        ctx = await browser.new_context(no_viewport=True)
        return ctx, browser, False


async def _arun(job: dict) -> None:
    import asyncio

    from playwright.async_api import async_playwright

    url = job.get("url")
    if not url or url == "#":
        print("This job has no application URL.")
        return
    vals = _profile_values()
    panel_data = {"title": job.get("title", ""), "company": job.get("company", ""),
                  "match": _skill_match(job)}
    print(f"\n▶ Opening application for: {job.get('title', '')} — {job.get('company', '')}")
    print(f"  {url}\n")

    async with async_playwright() as pw:
        ctx, _browser, attached = await _launch(pw)
        # Attached to your Chrome → open a fresh tab (don't hijack an existing one).
        page = await ctx.new_page() if attached else (
            ctx.pages[0] if ctx.pages else await ctx.new_page())

        # ── Panel actions, wired to the injected buttons (async = no re-entrancy deadlock) ──
        async def _on_autofill(source):
            log = await _afill(page, vals)
            if job.get("id"):
                _mark_drafted(job["id"])
            return {"count": len(log), "log": log}

        async def _on_cover(source):
            # claude CLI is blocking — run it off the event loop so the UI stays live.
            return await asyncio.get_event_loop().run_in_executor(None, _cover_letter, job)

        async def _on_cover_docx(source, body):
            def build():
                path = _build_cover_docx(job, body or "")
                _open_file(path)
                import os
                return os.path.basename(path)
            return await asyncio.get_event_loop().run_in_executor(None, build)

        async def _on_insert(source, text):
            return await _afill_cover_field(page, text or "")

        async def _on_applied(source):
            if job.get("id"):
                _mark_applied(job["id"])
            return True

        async def _on_questions(source):
            data = await page.evaluate(_JS_COLLECT)
            out = []
            for f in data["fields"]:
                is_text = f["tag"] == "textarea" or (f["tag"] == "input" and f["type"] in ("text", ""))
                if is_text and _answerable(f["label"]):
                    out.append({"idx": f["idx"], "label": f["label"]})
            return out

        async def _on_answer(source, question, hint):
            # claude CLI is blocking — off the event loop so the panel stays responsive.
            return await asyncio.get_event_loop().run_in_executor(
                None, _answer_question, job, question or "", hint or "")

        async def _on_insert_field(source, idx, text):
            try:
                return bool(await page.evaluate(_INSERT_JS, {"idx": idx, "text": text or ""}))
            except Exception:
                return False

        async def _on_experience(source):
            return _work_experience()

        await page.expose_binding("jhAutofill", _on_autofill)
        await page.expose_binding("jhCover", _on_cover)
        await page.expose_binding("jhCoverDocx", _on_cover_docx)
        await page.expose_binding("jhInsertCover", _on_insert)
        await page.expose_binding("jhApplied", _on_applied)
        await page.expose_binding("jhQuestions", _on_questions)
        await page.expose_binding("jhAnswer", _on_answer)
        await page.expose_binding("jhInsertField", _on_insert_field)
        await page.expose_binding("jhExperience", _on_experience)

        async def _inject():
            try:
                await page.evaluate(_PANEL_JS, panel_data)
            except Exception:
                pass

        page.on("load", lambda *a: asyncio.create_task(_inject()))   # re-inject after navigations

        await page.goto(url, wait_until="domcontentloaded", timeout=45000)
        await page.wait_for_timeout(2500)            # let the form render

        # Some ATS hide the form behind an "Apply" button — click it if present.
        for txt in ("Apply for this job", "Apply", "I'm interested", "Apply Now"):
            try:
                btn = page.get_by_role("button", name=txt, exact=False).first
                if await btn.is_visible(timeout=800):
                    await btn.click(); await page.wait_for_timeout(1500); break
            except Exception:
                continue

        await _inject()
        print("▶ Panel loaded — use it to autofill, draft a cover letter, and mark applied.")
        print("  ⚠ REVIEW EVERY FIELD before you submit. Close the window when done.\n")

        try:
            while True:
                await page.wait_for_timeout(500)     # alive until you close the window
        except Exception:
            pass


def run(job: dict) -> None:
    import asyncio
    asyncio.run(_arun(job))


def _mark_applied(job_id: int) -> None:
    """Mark applied + log a submitted application — mirrors the web /applied route
    and stamps the tracker lane so the card lands in Applied."""
    try:
        with db.connect() as conn:
            conn.execute("UPDATE jobs SET status='applied', track_status='applied' WHERE id=?", (job_id,))
            n = conn.execute(
                "UPDATE applications SET status='submitted', submitted_at=CURRENT_TIMESTAMP "
                "WHERE job_id=? AND status='draft'", (job_id,)).rowcount
            has = conn.execute("SELECT 1 FROM applications WHERE job_id=? AND status='submitted'",
                               (job_id,)).fetchone()
            if not n and not has:
                conn.execute("INSERT INTO applications (job_id, status, submitted_at) "
                             "VALUES (?, 'submitted', CURRENT_TIMESTAMP)", (job_id,))
            conn.commit()
    except Exception:
        pass


def _mark_drafted(job_id: int) -> None:
    try:
        with db.connect() as conn:
            conn.execute("UPDATE jobs SET status='drafted' WHERE id=? AND status IN ('new','')", (job_id,))
            conn.execute(
                "INSERT INTO applications (job_id, status) VALUES (?, 'draft')", (job_id,))
            conn.commit()
    except Exception:
        pass


def _job_by_id(job_id: int) -> dict | None:
    with db.connect() as conn:
        r = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
        return dict(r) if r else None


def main(argv: list[str]) -> int:
    if not argv:
        print("usage: python -m jobhunt.apply.autofill <job_id> | --url <url>")
        return 2
    if argv[0] == "--url":
        run({"url": argv[1], "title": "(ad-hoc)", "company": ""})
        return 0
    job = _job_by_id(int(argv[0]))
    if not job:
        print(f"No job with id {argv[0]}")
        return 1
    run(job)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
